# -*- coding: utf-8 -*-
"""
Created on Fri Oct 30 00:08:26 2020

@author: Ankit Parashar
"""
#! python3
# createInvites.py
import docx

inviteDoc = docx.Document()
guestFile = open('guests.txt', 'r')
guests = guestFile.readlines()
noOfGuests = len(guests)
for i in range(noOfGuests):
    inviteDoc.add_paragraph('It would be a pleasure to have the company of\n', 'Heading 1')
    inviteDoc.add_paragraph(guests[i].strip()+'\n', 'Heading 2')
    inviteDoc.add_paragraph('at 11010 Memory Lane on the Evening of\n', 'Heading 1')
    inviteDoc.add_paragraph('April 1\n', 'Heading 3')
    inviteDoc.add_paragraph('at 7\'o clock\n', 'Heading 1')
    if i != noOfGuests - 1:
        inviteDoc.paragraphs[4+5*i].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
inviteDoc.save('invitations.docx')